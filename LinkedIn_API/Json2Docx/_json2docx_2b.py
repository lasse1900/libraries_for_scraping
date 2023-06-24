import json
from docx import Document

# This one -----------------------//----------------------------

# Load JSON data from the file
# without title data_b.json
with open('sorted_2_out.json', 'r') as json_file: 
# with title    
# with open('data.json', 'r') as json_file: 
    json_data = json.load(json_file)

# Create a new Word document
document = Document()

# Iterate through JSON data and add content to the document
for item in json_data:
    # Add the document title
    # title = item['title']
    # title = "title"
    # document.add_heading(title, level=1)

    # Add sections to the document
    data = item['data']
    for section in data:
        company_name = section['job_location']
        date = section['date']
        job_location = section['job_location']
        job_url = section['job_url']

        # Add section job_location
        document.add_heading(company_name, level=2)

        # Add section job_location
        document.add_paragraph(job_url)

        # Add section job_location
        document.add_paragraph(date)

        # Add section job_location
        document.add_paragraph(job_location)

# Save the Word document
document.save('output_1505.docx')
