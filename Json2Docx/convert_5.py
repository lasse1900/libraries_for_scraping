import docx

doc = docx.Document()

doc.add_paragraph("This is some text")
paraObject = doc.add_paragraph("I love Python")
paraObject.add_run(' I am continuing this paragraph')

doc.save('new.docx')