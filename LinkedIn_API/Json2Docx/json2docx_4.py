import json
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import nsdecls

# Load JSON data from the file
with open('data.json', 'r') as json_file:
    json_data = json.load(json_file)

# Create a new Word document
document = Document()

# Iterate through JSON data and add content to the document
for item in json_data:
    title = item['title']
    document.add_heading(title, level=1)

    sections = item['sections']
    for section in sections:
        heading = section['heading']
        content = section['content']
        hyperlink = section['hyperlink']

        # Add section heading
        document.add_heading(heading, level=2)

        # Create hyperlink element
        hyperlink_element = document.add_paragraph().add_run().add_hyperlink(hyperlink, heading)

        # Set hyperlink color to blue
        hyperlink_element.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK

        # Add section content
        document.add_paragraph(content)

# Save the Word document
document.save('output_1405c.docx')
