from docx import Document
import json

doc = Document()
doc.add_heading('Open Jobs', 0)

with open('sorted_with_header.json', encoding='utf-8') as user_file:
    file_contents = json.load(user_file)

print(type(file_contents)) # dict

for job in file_contents['Jobs']:
    doc.add_paragraph(job['company_name'])
    doc.add_paragraph(job['job_location'])
    doc.add_paragraph(job['date'])
    print(job['job_url'])
    doc.add_paragraph(job['job_url'])

    doc.save('jobs.docx')