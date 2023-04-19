# pip install docx
# pip install document
# pip install python-docx
# pip install Path
# pip install pathlib

import re
from pathlib import Path

from docx import Document

path = Path(r"c:\coding\text.txt")

if path.exists():
    print(path, "exists")
else:
    print(path, "does not exist")
    raise SystemExit(-1)


for file in path.glob("*"):
    # file is a Path object

    document = Document()
    # file.name is the name of the file as str without the Path
    document.add_heading(file.name, 0)

    # Path objects do have the read_text, read_bytes
    # method and also supports
    # open with context managers

    # remove all non-XML-compatible characters
    file_content = re.sub(r"[^\x00-\x7F]+|\x0c", " ", file.read_text())
    document.add_paragraph(file_content)
    # if Document could not handle Path objects,
    # you must convert the Path object to a str

    document.save(str(file))
    document.save(file)