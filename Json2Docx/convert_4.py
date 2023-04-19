import docx
import os
# https://www.youtube.com/watch?v=26vNgM_wSAE&t=158s

os.getcwd()

doc = docx.Document('python.docx')
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)

print(doc.paragraphs[0].runs[0].text)

print(doc.paragraphs[0].runs)

print(doc.paragraphs[0].runs[1].text)
print(doc.paragraphs[0].runs[3].text)
print("-"*20)
print(doc.paragraphs[2].runs[0].text)
print(doc.paragraphs[3].runs[0].text)

