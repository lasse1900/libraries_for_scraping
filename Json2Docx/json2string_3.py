from docx import Document
from docx.shared import Inches
import json

doc = Document()
doc.add_heading('Open Jobs', 0)

with open('sorted_header.json') as user_file:
    file_contents = json.load(user_file)

print(type(file_contents)) # dict

for job in file_contents['Jobs']:
    doc.add_paragraph(job['company_name'])
    doc.add_paragraph(job['job_location'])
    doc.add_paragraph(job['date'])
    print(job['job_url'])
    doc.add_paragraph(job['job_url'])

    doc.save('jobs.docx')