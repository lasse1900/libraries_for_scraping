from xmlrpc.server import DocXMLRPCRequestHandler
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Create a new document
doc = Document()

# Add a paragraph with a hyperlink
paragraph = doc.add_paragraph()
hyperlink = paragraph.add_run("Example Hyperlink")
hyperlink.font.size = Pt(12)
hyperlink.font.underline = True

# Set the hyperlink relationship
r = hyperlink._element
r.rPr.rStyle = None
hyperlink._hyperlink = parse_xml(
    r'<w:hyperlink %s r:id="rId1" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:r><w:rPr><w:rStyle w:val="Hyperlink" /></w:rPr><w:t>Example Hyperlink</w:t></w:r></w:hyperlink>' % nsdecls('r')
)

# Set the relationship ID
relationship = doc.part.relate_to(
    "http://www.example.com",
    DocXMLRPCRequestHandler.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
    is_external=True,
)
hyperlink._hyperlink.rId = relationship.rId

# Save the document
doc.save("output.docx")
