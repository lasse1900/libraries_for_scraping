import json
from docx import Document

# Open the JSON file for reading
with open('filtered.json', 'r') as file:
    json_data = json.load(file)

# Add the desired clause at the beginning
modified_data = [{'data': json_data}]

# Open the same JSON file for writing (which overwrites the existing file)
with open('filtered_with_header.json', 'w') as file:
    json.dump(modified_data, file, indent=4)

# Load JSON data from the file and add a JSON header
with open('filtered_with_header.json', 'r') as json_file: 
    json_data = json.load(json_file)

# Create a new Word document
document = Document()

# Iterate through JSON data and add content to the document
for item in json_data:
    # Add the document title
    title = "Open jobs on fi.Indeed.com"
    document.add_heading(title, level=1)

    # Add sections to the document
    data = item['data']
    for section in data:
        company_name = section['company_name']
        job_title = section['job_title']
        date = section['date']
        job_location = section['job_location']
        job_url = section['job_url']

        # Add sections to docx
        document.add_heading(company_name, level=2)
        document.add_paragraph(job_title)        
        document.add_paragraph(job_url)
        document.add_paragraph(date)
        document.add_paragraph(job_location)
        document.add_paragraph('_'*50)

document.save('jobs.docx')
