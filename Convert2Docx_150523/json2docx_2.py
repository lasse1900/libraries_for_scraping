import json
from docx import Document

# Load JSON data from the file and add a JSON header
with open('filtered_with_header.json', 'r') as json_file: 
    json_data = json.load(json_file)

# Create a new Word document
document = Document()

# Iterate through JSON data and add content to the document
for item in json_data:
    # Add the document title
    title = "Open jobs on fi.Indeed.com"
    document.add_heading(title, level=1)

    # Add sections to the document
    data = item['data']
    for section in data:
        company_name = section['company_name']
        job_title = section['job_title']
        date = section['date']
        job_location = section['job_location']
        job_url = section['job_url']

        # Add section job_location
        document.add_heading(company_name, level=2)
        # Add section job_location
        document.add_paragraph(job_title)        
        # Add section job_location
        document.add_paragraph(job_url)
        # Add section job_location
        document.add_paragraph(date)
        # Add section job_location
        document.add_paragraph(job_location)
        document.add_paragraph('-'*110)

# Save the Word document
document.save('jobs.docx')
