import json
from docx import Document

# Load JSON data from the file
# without title data_b.json
with open('data_b.json', 'r') as json_file: 
# with title    
# with open('data.json', 'r') as json_file: 
    json_data = json.load(json_file)

# Create a new Word document
document = Document()

# Iterate through JSON data and add content to the document
for item in json_data:
    # Add the document title
    # title = item['title']
    # document.add_heading(title, level=1)

    # Add sections to the document
    sections = item['sections']
    for section in sections:
        heading = section['heading']
        content = section['content']

        # Add section heading
        document.add_heading(heading, level=2)

        # Add section content
        document.add_paragraph(content)

# Save the Word document
document.save('output_1505b.docx')
