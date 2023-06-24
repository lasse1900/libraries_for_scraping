import json
from docx import Document

# Load JSON data from the file
with open('data.json', 'r') as json_file:
    json_data = json.load(json_file)

# Create a new Word document
document = Document()

# Iterate through JSON data and add content to the document
for item in json_data:

    # Add sections to the document
    data = item['sections']
    for item in data:
        heading = item['heading']
        content = item['content']

        # Add section heading
        document.add_heading(heading, level=2)

        # Add section content
        document.add_paragraph(content)

# Save the Word document
document.save('output_1405b.docx')
