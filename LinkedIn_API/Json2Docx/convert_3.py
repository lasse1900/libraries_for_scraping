import re
import os
from pathlib import Path
import sys
import docx

# The location where the files are located
input_path = r'c:\coding\input'
# Locatia unde vom scrie fisierele docx
output_path = r'c:\coding\output'
# The location where we will write the docx files
os.makedirs(output_path, exist_ok=True)
 
# Check the existence of the folder
directory_path = Path(input_path)
if directory_path.exists() and directory_path.is_dir():
    print(directory_path, "exists")
else:
    print(directory_path, "is invalid")
    sys.exit(1)
 
for file_path in directory_path.glob("*"):
    # file_path is a Path object
 
    print("Processing the file:", file_path)
    document = docx.Document()
    # file_path.name is the name of the file as str without the Path
    document.add_heading(file_path.name, 0)
 
    file_content = file_path.read_text(encoding='UTF-8')
    document.add_paragraph(file_content)
 
    # build the new path where we store the files
    output_file_path = os.path.join(output_path, file_path.name + ".docx")
 
    document.save(output_file_path)
    print("Converted the following file:", file_path, "in: ", output_file_path)