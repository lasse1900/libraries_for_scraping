import json
from xmlrpc.server import DocXMLRPCRequestHandler, DocXMLRPCServer
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url,DocXMLRPCServer.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = DocXMLRPCRequestHandler.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(nsdecls('r'), 'id', r_id)
    hyperlink.append(DocXMLRPCRequestHandler.oxml.shared.OxmlElement('w:r'))
    new_run = hyperlink.getchildren()[0]

    run = paragraph.add_run()
    run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    run.font.underline = True

    new_run.append(run._r)
    run.text = text
    paragraph._p.insert(0, hyperlink)

# Load JSON data from file
with open('data_b.json') as json_file:
    data = json.load(json_file)

# Create a new Word document
document = Document()

# Iterate through JSON data and add content to the document
for item in data:
    print("round", item)
    # Add the document title
    # title = item['title']
    # document.add_heading(title, level=1)

    # Add sections to the document
    sections = item['sections']
    for section in sections:
        print("section", section)
        heading = section['heading']
        content = section['content']
        hyperlink = section['hyperlink']


    # Add a heading
    # doc.add_heading(title, level=1)

    # Add content
    paragraph = document.add_paragraph()
    paragraph.add_run(content)

    # Add hyperlink
    add_hyperlink(paragraph, hyperlink, hyperlink)
    document.add_paragraph(content)

# Save the document
# doc.save('output.docx')
document.save('output_1505b.docx')
