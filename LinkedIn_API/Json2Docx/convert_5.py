import docx

doc = docx.Document()

job_url = "https://fi.indeed.com/viewjob?viewtype=embedded&jk=5d2119e1f4d929b4&from=vjs&tk=1gtqnlvtsir0k800&continueUrl=%2Fjobs%3Ffilter%3D0%26q%3D%2527django%2527%26start%3D0%26l%3D%2527tampere%2527"

doc.add_paragraph("This is some text")
paraObject = doc.add_paragraph("I love Python")
paraObject.add_run(' I am continuing this paragraph')
paraObject = doc.add_paragraph("Ends here ...")
doc.add_paragraph("https://pyfpdf.readthedocs.io/en/latest/index.html")

doc.save('new2.docx')