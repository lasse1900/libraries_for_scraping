
from docx import Document

document = Document()
document.add_heading('Open Jobs listing with following keywords:')
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')

document.add_page_break()
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
document.save('test.docx')
